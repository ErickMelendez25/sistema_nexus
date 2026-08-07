"""
routers/fill_document.py
========================
Endpoints para auto-rellenar documentos DOCX/PDF
con datos de la empresa desde la BD.

POST /api/fill-document          → detecta campos y devuelve preview editable
POST /api/fill-document/generar  → genera el DOCX final con valores confirmados
GET  /api/empresas               → lista empresas disponibles
"""

import os
import re
import io
import logging
import tempfile
import shutil
from pathlib import Path
from datetime import datetime
from typing import Optional, List, Dict, Any

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import mysql.connector
from mysql.connector import pooling
from docx import Document
from docxtpl import DocxTemplate

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api", tags=["fill-document"])

# ─── STORAGE (mismo que contratos.py) ────────────────────────────────────────
STORAGE_BACKEND = os.getenv("STORAGE_BACKEND", "local")
AZURE_CONN_STR  = os.getenv("AZURE_STORAGE_CONNECTION_STRING", "")
AZURE_CONTAINER = os.getenv("AZURE_CONTAINER_NAME", "seace-archivos")
AWS_BUCKET      = os.getenv("AWS_BUCKET_NAME", "seace-archivos")
AWS_REGION      = os.getenv("AWS_REGION", "us-east-1")

_DOCS = Path("D:/")
if not _DOCS.exists():
    _DOCS = Path.home()
CARPETA_SALIDA_LOCAL = _DOCS / "SEACE_PLADIBOT"

# ─── POOL BD ──────────────────────────────────────────────────────────────────
_pool: Optional[pooling.MySQLConnectionPool] = None

def get_pool():
    global _pool
    if _pool is None:
        _pool = pooling.MySQLConnectionPool(
            pool_name="fill_doc",
            pool_size=3,
            host     = os.getenv("MYSQL_HOST",     "localhost"),
            user     = os.getenv("MYSQL_USER",     "root"),
            password = os.getenv("MYSQL_PASSWORD", "Erick2026#"),
            database = os.getenv("MYSQL_DATABASE", "pladibot_db"),
            charset  = "utf8mb4",
        )
    return _pool

def get_conn():
    return get_pool().get_connection()

# ─── CATÁLOGO DE CAMPOS ───────────────────────────────────────────────────────
# Mapea patrones que aparecen en documentos SEACE → clave interna
CAMPO_PATRONES: List[tuple] = [
    # RUC — debe ir ANTES que dirección/domicilio
    (r'n[°º]?\s*ruc\s*[:：]',           'empresa_ruc'),
    (r'^\s*ruc\s*[:：]',                'empresa_ruc'),
    (r'\bruc\b',                        'empresa_ruc'),
    # Razón social / Nombre empresa
    (r'raz[oó]n\s+social',              'empresa_razon_social'),
    (r'nombre[,\s]+denominaci[oó]n\s+o\s+raz[oó]n',  'empresa_razon_social'),
    (r'nombre\s+de\s+la\s+empresa',     'empresa_razon_social'),
    (r'empresa\s+postora',              'empresa_razon_social'),
    (r'nombre\s+y/o\s+raz[oó]n\s+social', 'empresa_razon_social'),
    # DNI — debe ir ANTES que domicilio
    (r'dni\s*n[°º]?\s*[:：]',           'representante_dni'),
    (r'^\s*dni\s*[:：]',                'representante_dni'),
    (r'\bdni\b',                        'representante_dni'),
    (r'documento\s+de\s+identidad',     'representante_dni'),
    # Dirección / Domicilio
    (r'domicilio\s+legal\s*[:：]',      'empresa_direccion'),
    (r'direcci[oó]n\s+legal\s*[:：]',   'empresa_direccion'),
    (r'domiciliado\s+en\s*[:：]',       'empresa_direccion'),
    # Teléfono
    (r'tel[eé]fono\s*[:（(]',           'empresa_telefono'),
    (r'tel[eé]fono\s*[:：]',            'empresa_telefono'),
    (r'\btelefono\b',                   'empresa_telefono'),
    (r'celular',                        'empresa_telefono'),
    # Correo
    (r'correo\s+electr[oó]nico\s*[:：]','empresa_correo'),
    (r'e[\-\s]?mail\s*[:：]',           'empresa_correo'),
    (r'correo\s+electr[oó]nico',        'empresa_correo'),
    # Representante
    (r'representante\s+legal',          'representante_nombre'),
    (r'firma\s+y/o\s+sello',            'representante_nombre'),
    (r'nombre\s+y/o\s+raz[oó]n',        'representante_nombre'),
    (r'persona\s+de\s+contacto',        'representante_nombre'),
    (r'nombres\s+y\s+apellidos',        'representante_nombre'),
    # Cargo
    (r'\bcargo\b',                      'representante_cargo'),
    # Fecha — patrones específicos de docs SEACE
    (r'lima\s*,\s*\w*\s+de\s+\w*\s+de\s+20\d\d', 'fecha_actual'),
    (r'lugar\s+y\s+fecha',              'fecha_actual'),
    (r'\bfecha\b',                      'fecha_actual'),
]

# Placeholders visuales que reemplazamos
PLACEHOLDER_RE = re.compile(
    r'_{3,}|\.{3,}|<[^>]{1,50}>|\[[^\]]{1,50}\]|\({3,}\)'
)

def detectar_campo(texto: str) -> Optional[str]:
    """Dado el texto de un párrafo/celda, devuelve la clave del campo detectado."""
    t = texto.lower().strip()
    for patron, campo in CAMPO_PATRONES:
        if re.search(patron, t, re.IGNORECASE):
            return campo
    return None

def obtener_empresa(empresa_id: int) -> Dict[str, Any]:
    conn = get_conn()
    cur  = conn.cursor(dictionary=True)
    try:
        cur.execute("SELECT * FROM empresas WHERE id = %s", (empresa_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(404, f"Empresa {empresa_id} no encontrada")
        return row
    finally:
        cur.close()
        conn.close()

def valores_empresa(empresa: Dict) -> Dict[str, str]:
    """Construye el mapa campo → valor para rellenar."""
    hoy = datetime.now().strftime("%d/%m/%Y")
    return {
        'empresa_ruc'         : empresa.get('ruc', ''),
        'empresa_razon_social': empresa.get('razon_social', ''),
        'empresa_direccion'   : empresa.get('direccion', ''),
        'empresa_telefono'    : empresa.get('telefono', ''),
        'empresa_correo'      : empresa.get('correo', ''),
        'representante_nombre': empresa.get('representante_nombre', ''),
        'representante_dni'   : empresa.get('representante_dni', ''),
        'representante_cargo' : empresa.get('representante_cargo', ''),
        'fecha_actual'        : hoy,
    }

# ─── OBTENER RUTA LOCAL DEL ARCHIVO (igual lógica que contratos.py) ───────────
def resolver_ruta(ruta_local: str) -> Optional[Path]:
    if not ruta_local:
        return None

    # Azure/AWS: ruta_local es una URL
    if ruta_local.startswith("http"):
        return None  # se maneja aparte

    ruta = Path(ruta_local)

    # ── Caso 1: ruta absoluta que existe tal cual ─────────────────────────
    if ruta.exists():
        return ruta

    # ── Caso 2: ruta relativa (ej: archivos\73597\cotizacion\287918_...)
    #    Construir desde CARPETA_SALIDA_LOCAL (Documents\SEACE_PLADIBOT)
    if not ruta.is_absolute():
        ruta2 = CARPETA_SALIDA_LOCAL / ruta
        if ruta2.exists():
            return ruta2

    # ── Caso 3: ruta absoluta de otra PC — auto-repair buscando SEACE_PLADIBOT
    partes = ruta.parts
    try:
        idx      = next(i for i, p in enumerate(partes) if p == "SEACE_PLADIBOT")
        relativa = Path(*partes[idx + 1:])   # todo lo que va DESPUÉS de SEACE_PLADIBOT
        ruta3    = CARPETA_SALIDA_LOCAL / relativa
        if ruta3.exists():
            return ruta3
    except StopIteration:
        pass

    # ── Caso 4: ruta absoluta vieja con "seace_output" (rutas pre-migración)
    partes = ruta.parts
    try:
        idx      = next(i for i, p in enumerate(partes) if p == "seace_output")
        relativa = Path(*partes[idx + 1:])   # todo lo que va DESPUÉS de seace_output
        ruta4    = CARPETA_SALIDA_LOCAL / relativa
        if ruta4.exists():
            return ruta4
    except StopIteration:
        pass

    return None

def descargar_bytes_nube(ruta_url: str, extension: str) -> Optional[bytes]:
    """Descarga archivo desde Azure o AWS a memoria."""
    if STORAGE_BACKEND == "azure":
        try:
            from azure.storage.blob import BlobServiceClient
            blob_name = ruta_url.split("/")[-1].split("?")[0]
            client    = BlobServiceClient.from_connection_string(AZURE_CONN_STR)
            blob      = client.get_blob_client(container=AZURE_CONTAINER, blob=blob_name)
            return blob.download_blob().readall()
        except Exception as e:
            logger.error(f"Azure descarga error: {e}")
            return None

    if STORAGE_BACKEND == "aws":
        try:
            import boto3
            s3_key = ruta_url.split("/")[-1].split("?")[0]
            s3     = boto3.client("s3", region_name=AWS_REGION)
            obj    = s3.get_object(Bucket=AWS_BUCKET, Key=s3_key)
            return obj["Body"].read()
        except Exception as e:
            logger.error(f"AWS descarga error: {e}")
            return None

    return None

# ─── MODELOS ──────────────────────────────────────────────────────────────────

class CampoDetectado(BaseModel):
    indice      : int
    campo       : str           # clave interna ej: empresa_ruc
    label       : str           # label legible ej: "RUC"
    valor_auto  : str           # valor que vamos a poner
    editable    : bool = True
    contexto    : str  = ""     # texto original del párrafo donde se detectó

class PreviewFill(BaseModel):
    id_archivo  : int
    nombre      : str
    extension   : str
    campos      : List[CampoDetectado]
    tiene_campos: bool

class FillRequest(BaseModel):
    id_contrato : int
    id_archivo  : int
    empresa_id  : int

class GenerarRequest(BaseModel):
    id_contrato : int
    id_archivo  : int
    empresa_id  : int
    valores     : Dict[str, str]   # campo → valor corregido por usuario

LABELS = {
    'empresa_ruc'         : 'RUC',
    'empresa_razon_social': 'Razón Social',
    'empresa_direccion'   : 'Dirección',
    'empresa_telefono'    : 'Teléfono',
    'empresa_correo'      : 'Correo electrónico',
    'representante_nombre': 'Representante Legal',
    'representante_dni'   : 'DNI del Representante',
    'representante_cargo' : 'Cargo del Representante',
    'fecha_actual'        : 'Fecha',
}

# ─── ENDPOINTS ────────────────────────────────────────────────────────────────

@router.get("/empresas")
def listar_empresas():
    """Lista empresas disponibles para el selector del frontend."""
    conn = get_conn()
    cur  = conn.cursor(dictionary=True)
    try:
        cur.execute("""
            SELECT id, razon_social, ruc, representante_nombre
            FROM empresas WHERE activa = 1 ORDER BY razon_social
        """)
        return cur.fetchall()
    finally:
        cur.close()
        conn.close()


@router.post("/fill-document", response_model=PreviewFill)
def detectar_campos(req: FillRequest):
    """
    Detecta los campos rellenables en el DOCX/PDF y devuelve
    un preview con los valores auto-detectados listos para editar.
    """
    # 1. Buscar archivo en BD
    conn = get_conn()
    cur  = conn.cursor(dictionary=True)
    try:
        cur.execute("""
            SELECT id_archivo, nombre, extension, ruta_local
            FROM archivos
            WHERE id_contrato = %s AND id_archivo = %s
            LIMIT 1
        """, (req.id_contrato, req.id_archivo))
        arch = cur.fetchone()
    finally:
        cur.close()
        conn.close()

    if not arch:
        raise HTTPException(404, "Archivo no encontrado en BD")

    extension = (arch.get('extension') or '').lower().strip().lstrip('.')
    logger.info(f"[fill] extension='{extension}' archivo='{arch.get('nombre')}'")
    if extension not in ('docx', 'doc', 'pdf'):
        raise HTTPException(400, f"Formato no soportado: {extension}")

    # 2. Obtener empresa y valores
    empresa = obtener_empresa(req.empresa_id)
    vals    = valores_empresa(empresa)

    # 3. Obtener bytes del archivo
    ruta_local = arch.get('ruta_local', '')
    contenido  = None

    if ruta_local and ruta_local.startswith("http"):
        contenido = descargar_bytes_nube(ruta_local, extension)
    else:
        ruta = resolver_ruta(ruta_local)
        if ruta:
            contenido = ruta.read_bytes()

    if not contenido:
        raise HTTPException(404, f"Archivo no encontrado en disco/nube: {ruta_local}")

    # 4. Leer DOCX y detectar campos
    # 4. Leer documento y detectar campos
    campos = []
    indice = 0

    def analizar_texto(texto: str):
        nonlocal indice
        campo = detectar_campo(texto)
        if campo and campo not in [c.campo for c in campos]:
            campos.append(CampoDetectado(
                indice     = indice,
                campo      = campo,
                label      = LABELS.get(campo, campo),
                valor_auto = vals.get(campo, ''),
                editable   = True,
                contexto   = texto.strip()[:120],
            ))
            indice += 1

    if extension == 'pdf':
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(contenido)) as pdf:
                for page in pdf.pages:
                    texto_pagina = page.extract_text() or ''
                    for linea in texto_pagina.split('\n'):
                        if linea.strip():
                            analizar_texto(linea)
        except ImportError:
            raise HTTPException(500, "Instala pdfplumber: pip install pdfplumber")

    elif extension == 'doc':
        texto_doc = ''

        # Intento 1: mammoth
        try:
            import mammoth
            resultado = mammoth.extract_raw_text(io.BytesIO(contenido))
            texto_doc = resultado.value or ''
        except Exception as e_mammoth:
            logger.warning(f"mammoth falló con DOC: {e_mammoth}")

        # Intento 2: oletools (msodumper) si mammoth falló
        if not texto_doc.strip():
            try:
                from olefile import OleFileIO
                ole = OleFileIO(io.BytesIO(contenido))
                if ole.exists('WordDocument'):
                    # Extraer texto crudo del stream de Word
                    stream = ole.openstream('WordDocument').read()
                    # Texto está en el stream de tablas de strings
                    texto_doc = stream.decode('latin-1', errors='ignore')
                    # Limpiar caracteres no imprimibles
                    texto_doc = re.sub(r'[^\x20-\x7E\xA0-\xFF\n\r\t]', ' ', texto_doc)
                ole.close()
            except Exception as e_ole:
                logger.warning(f"olefile falló con DOC: {e_ole}")

        # Intento 3: leer como binario y extraer strings legibles
        if not texto_doc.strip():
            try:
                texto_doc = contenido.decode('latin-1', errors='ignore')
                texto_doc = re.sub(r'[^\x20-\x7E\xA0-\xFF\n\r\t]', ' ', texto_doc)
                # Quedarse solo con palabras de más de 3 chars
                palabras = re.findall(r'[A-Za-záéíóúÁÉÍÓÚñÑ]{4,}[\s\w]*', texto_doc)
                texto_doc = '\n'.join(palabras)
                logger.info(f"[fill] DOC leído como binario crudo, {len(palabras)} palabras")
            except Exception as e_raw:
                raise HTTPException(500, f"No se pudo leer el DOC por ningún método: mammoth={e_mammoth}, raw={e_raw}")

        if not texto_doc.strip():
            raise HTTPException(500, "El archivo DOC no tiene texto extraíble")

        for linea in texto_doc.split('\n'):
            if linea.strip():
                analizar_texto(linea)

    else:
        # DOCX
        try:
            doc = Document(io.BytesIO(contenido))
            for para in doc.paragraphs:
                analizar_texto(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        analizar_texto(cell.text)
        except Exception as e:
            # Fallback: intentar con mammoth si python-docx falla
            try:
                import mammoth
                resultado = mammoth.extract_raw_text(io.BytesIO(contenido))
                for linea in resultado.value.split('\n'):
                    if linea.strip():
                        analizar_texto(linea)
            except Exception as e2:
                raise HTTPException(500, f"Error leyendo DOCX: {e} | fallback: {e2}")

    return PreviewFill(
        id_archivo  = arch['id_archivo'],
        nombre      = arch.get('nombre', ''),
        extension   = extension,
        campos      = campos,
        tiene_campos= len(campos) > 0,
    )


@router.post("/fill-document/generar")
def generar_documento(req: GenerarRequest):
    """
    Genera el DOCX final con los valores confirmados por el usuario
    y lo devuelve como descarga directa.
    """
    # 1. Buscar archivo en BD
    conn = get_conn()
    cur  = conn.cursor(dictionary=True)
    try:
        cur.execute("""
            SELECT nombre, extension, ruta_local
            FROM archivos
            WHERE id_contrato = %s AND id_archivo = %s
            LIMIT 1
        """, (req.id_contrato, req.id_archivo))
        arch = cur.fetchone()
    finally:
        cur.close()
        conn.close()

    if not arch:
        raise HTTPException(404, "Archivo no encontrado")

    extension  = (arch.get('extension') or '').lower().strip('.')

    ruta_local = arch.get('ruta_local', '')

    # 2. Obtener bytes
    contenido = None
    if ruta_local and ruta_local.startswith("http"):
        contenido = descargar_bytes_nube(ruta_local, extension)
    else:
        ruta = resolver_ruta(ruta_local)
        if ruta:
            contenido = ruta.read_bytes()

    if not contenido:
        raise HTTPException(404, "Archivo no encontrado en disco/nube")

    # 3. Si es PDF → generar DOCX nuevo con los valores
    if extension == 'pdf':
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc_nuevo = DocxDocument()

        # Título del documento
        titulo = doc_nuevo.add_heading('Documento Rellenado', level=1)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc_nuevo.add_paragraph('')

        # Agregar cada campo con su valor
        for campo, valor in req.valores.items():
            label = LABELS.get(campo, campo)
            p = doc_nuevo.add_paragraph()
            run_label = p.add_run(f'{label}: ')
            run_label.bold = True
            run_label.font.size = Pt(11)
            run_valor = p.add_run(valor or '—')
            run_valor.font.size = Pt(11)

        doc_nuevo.add_paragraph('')
        doc_nuevo.add_paragraph(f'Generado el: {datetime.now().strftime("%d/%m/%Y %H:%M")}')

        output = io.BytesIO()
        doc_nuevo.save(output)
        output.seek(0)

        nombre_original = arch.get('nombre', 'documento')
        nombre_salida   = f"RELLENADO_{Path(nombre_original).stem}.docx"

        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{nombre_salida}"'}
        )

    # 3b. Si es DOCX → reemplazar campos directamente
    # 3b. Si es DOC o DOCX → procesar
    if extension == 'doc':
        # DOC binario no es editable con python-docx
        # Generamos un DOCX nuevo con los valores igual que PDF
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc_nuevo = Document()
        titulo = doc_nuevo.add_heading('Documento Rellenado', level=1)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc_nuevo.add_paragraph('')
        for campo, valor in req.valores.items():
            label = LABELS.get(campo, campo)
            p = doc_nuevo.add_paragraph()
            run_label = p.add_run(f'{label}: ')
            run_label.bold = True
            run_label.font.size = Pt(11)
            run_valor = p.add_run(valor or '—')
            run_valor.font.size = Pt(11)
        doc_nuevo.add_paragraph('')
        doc_nuevo.add_paragraph(f'Generado el: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        output = io.BytesIO()
        doc_nuevo.save(output)
        output.seek(0)
        nombre_salida = f"RELLENADO_{Path(arch.get('nombre','documento')).stem}.docx"
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{nombre_salida}"'}
        )

    doc = Document(io.BytesIO(contenido))

# Patrón de texto basura previo (direcciones mal rellenadas antes)
    BASURA_RE = re.compile(
        r'PJ\.\s*ORELLANA[^,\n]*|'         # texto basura específico
        r'_{3,}|\.{4,}|'                    # guiones/puntos de relleno
        r'<[^>]{1,80}>|'                    # <campo>
        r'\[[^\]]{1,80}\]|'                 # [campo]
        r'\({3,}\)|'                        # (((
        r'\[CONSIGNAR[^\]]*\]',             # [CONSIGNAR ...]
        re.IGNORECASE
    )

    def reemplazar_en_parrafo(para):
        texto_original = para.text
        texto = texto_original

        for campo, valor in req.valores.items():
            label  = LABELS.get(campo, campo)
            valor_ = valor or ''

            # 1. Reemplazar placeholders visuales solos
            texto = BASURA_RE.sub(valor_, texto)

            # 2. Después de label: → reemplazar lo que sigue
            texto = re.sub(
                rf'({re.escape(label)}\s*[:：Nº°n°]*\s*)([\w\s.,/-]{{3,80}})',
                lambda m: m.group(1) + valor_,
                texto, flags=re.IGNORECASE
            )

        # Fecha especial: "Lima, de de 2026" → "Lima, 19/06/2026"
        fecha_val = req.valores.get('fecha_actual', '')
        if fecha_val:
            texto = re.sub(
                r'(Lima\s*,\s*)(\w+\s+de\s+\w+\s+de\s+20\d\d)',
                rf'\g<1>{fecha_val}',
                texto, flags=re.IGNORECASE
            )
            texto = re.sub(
                r'(Lima\s*,\s*)\w*\s+de\s+\w*\s+de\s+20\d\d',
                rf'\g<1>{fecha_val}',
                texto, flags=re.IGNORECASE
            )
            # Fecha vacía "Lima,      de                  de 2026"
            texto = re.sub(
                r'(Lima\s*,\s*)\s+de\s+\s+de\s+20\d\d',
                rf'\g<1>{fecha_val}',
                texto, flags=re.IGNORECASE
            )

        if texto != texto_original:
            if para.runs:
                run = para.runs[0]
                for r in para.runs[1:]:
                    r.text = ''
                run.text = texto
            else:
                para.text = texto

    for para in doc.paragraphs:
        reemplazar_en_parrafo(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    reemplazar_en_parrafo(para)

    # 4. Guardar en memoria y devolver
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    nombre_original = arch.get('nombre', 'documento')
    nombre_salida   = f"RELLENADO_{nombre_original}"
    if not nombre_salida.endswith('.docx'):
        nombre_salida += '.docx'

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{nombre_salida}"'}
    )