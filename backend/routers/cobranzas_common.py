"""
Lógica compartida para Cobranzas (DOC PARA PAGO y CARTA NOTA DÉBITO).

Es la misma lógica que ya tenías en cobranzas.py (CobranzasFrame) y
cobra.py (CobraFrame), pero separada de la interfaz tkinter para que
los dos routers de FastAPI la puedan reutilizar sin duplicar código.

Requisitos en el servidor (la PC Windows con Word instalado):
    pip install fastapi python-multipart python-docx pywin32 pypdf
"""
import os
import shutil
from datetime import datetime
from pathlib import Path
from typing import Iterable, List

import platform
import subprocess
from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader, PdfWriter

# Misma carpeta de salida que ya usan tus apps de escritorio
CARPETA_SALIDA = Path.home() / "Documents" / "CARTAS_GENERADAS"
CARPETA_SALIDA.mkdir(parents=True, exist_ok=True)

# Carpeta donde deben vivir cobranza_1.docx, cobranza_2.docx, cobra1.docx, cobra2.docx
# en el servidor. Ajusta esta ruta a donde realmente las tengas.
CARPETA_PLANTILLAS = Path(__file__).resolve().parent / "plantillas"

print("========================================")
print("CARPETA DE PLANTILLAS:", CARPETA_PLANTILLAS)
print("cobra1 existe:", (CARPETA_PLANTILLAS / "cobra1.docx").exists())
print("cobra2 existe:", (CARPETA_PLANTILLAS / "cobra2.docx").exists())
print("========================================")

MESES_ES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]


def formatear_fecha_es(fecha_str: str) -> str:
    fecha_str = (fecha_str or "").strip()
    # Prueba ambos formatos: "%Y-%m-%d" es lo que manda <input type="date">
    # del formulario web; "%d/%m/%Y" se deja como respaldo por si algún
    # llamador viejo (ej. la app de escritorio) sigue mandando dd/mm/aaaa.
    for formato in ("%Y-%m-%d", "%d/%m/%Y"):
        try:
            fecha = datetime.strptime(fecha_str, formato)
            return f"{fecha.day} de {MESES_ES[fecha.month - 1]} de {fecha.year}"
        except ValueError:
            continue
    return fecha_str

def _replace_in_paragraph(paragraph, data: dict) -> None:
    full_text = paragraph.text
    new_text = full_text
    replaced = False
    for key, value in data.items():
        marcador = f"{{{{{key}}}}}"
        if marcador in new_text:
            new_text = new_text.replace(marcador, str(value))
            replaced = True
    if replaced:
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.text = new_text


def _replace_in_table(table, data: dict) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _replace_in_paragraph(p, data)
            for nested_table in cell.tables:
                _replace_in_table(nested_table, data)


def reemplazar_marcadores(doc: Document, data: dict) -> None:
    for p in doc.paragraphs:
        _replace_in_paragraph(p, data)
    for table in doc.tables:
        _replace_in_table(table, data)
    for section in doc.sections:
        header, footer = section.header, section.footer
        for p in header.paragraphs:
            _replace_in_paragraph(p, data)
        for table in header.tables:
            _replace_in_table(table, data)
        for p in footer.paragraphs:
            _replace_in_paragraph(p, data)
        for table in footer.tables:
            _replace_in_table(table, data)


def _delete_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def _set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = text


def get_all_paragraphs(doc: Document) -> list:
    paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paras.extend(cell.paragraphs)
                for nt in cell.tables:
                    for nrow in nt.rows:
                        for ncell in nrow.cells:
                            paras.extend(ncell.paragraphs)
    return paras


def eliminar_marcadores_de_bloque(doc: Document) -> None:
    for p in list(get_all_paragraphs(doc)):
        txt = p.text
        if "{{BLOQUE_OCAM_INICIO}}" in txt or "{{BLOQUE_OCAM_FIN}}" in txt:
            _delete_paragraph(p)
            continue
        if "{{BLOQUE_DETALLE_INICIO}}" in txt or "{{BLOQUE_DETALLE_FIN}}" in txt:
            _delete_paragraph(p)
            continue


def ajustar_si_ocam_vacio(doc: Document, ocam_valor: str) -> None:
    if ocam_valor.strip():
        return
    for p in list(get_all_paragraphs(doc)):
        txt = p.text.strip()
        if txt.startswith("d)") and "Orden de Compra (Perú Compras)" in txt:
            _delete_paragraph(p)
            continue
        if txt.startswith("e)") and "Validez del comprobante de pago" in txt:
            _set_paragraph_text(p, "d)    Validez del comprobante de pago")
            continue
        if txt.startswith("f)") and "Carta CCI" in txt:
            _set_paragraph_text(p, "e)    Carta CCI")
            continue
        if txt.startswith("g)") and "Carta de Garantía" in txt:
            _set_paragraph_text(p, "f)     Carta de Garantía")
            continue



def ajustar_referencia_ocam(doc: Document, ocam_valor: str) -> None:
    """
    Para la plantilla CARTA NOTA DÉBITO (cobra1.docx / cobra2.docx), el
    párrafo de referencia trae:
        a) Orden de Compra Electrónica {{OCAM}}; Orden Física N° {{OC}}, SIAF N° {{SF}}

    Si OCAM viene vacío, se elimina SOLO el fragmento
    "Orden de Compra Electrónica {{OCAM}}; " y queda:
        a) Orden Física N° {{OC}}, SIAF N° {{SF}}

    IMPORTANTE: debe llamarse ANTES de reemplazar_marcadores(), porque
    busca el marcador {{OCAM}} literal en el texto (si ya se reemplazó
    por vacío, no hay forma de saber dónde estaba el fragmento).
    """
    if ocam_valor.strip():
        return

    fragmentos_posibles = [
        "Orden de Compra Electrónica {{OCAM}}; ",
        "Orden de Compra Electrónica {{OCAM}};",
        "Orden de Compra Electrónica {{OCAM}} ",
        "Orden de Compra Electrónica {{OCAM}}",
    ]

    parrafos = get_all_paragraphs(doc)
    for section in doc.sections:
        parrafos.extend(section.header.paragraphs)
        parrafos.extend(section.footer.paragraphs)

    for p in parrafos:
        txt = p.text
        for fragmento in fragmentos_posibles:
            if fragmento in txt:
                _set_paragraph_text(p, txt.replace(fragmento, ""))
                break



def quitar_texto_si(doc: Document, texto: str, condicion: bool) -> None:
    """
    Si `condicion` es True, borra `texto` (substring literal) de
    cualquier párrafo donde aparezca — cuerpo, tablas, encabezados y
    pies de página — sin tocar el resto del párrafo. Se usa para el
    checkbox "cargo sellado": la plantilla siempre trae el texto fijo
    " (cargo sellado)", y esto lo quita solo si el usuario destildó la
    casilla en el formulario.
    """
    if not condicion:
        return
    parrafos = get_all_paragraphs(doc)
    for section in doc.sections:
        parrafos.extend(section.header.paragraphs)
        parrafos.extend(section.footer.paragraphs)
    for p in parrafos:
        if texto in p.text:
            _set_paragraph_text(p, p.text.replace(texto, ""))





def _crear_pdf_word_windows(ruta_docx: str, ruta_pdf: str) -> None:
    """Convierte un .docx a .pdf usando Microsoft Word instalado (win32com)."""
    import pywintypes
    import win32com.client

    ruta_docx_abs = str(Path(ruta_docx).resolve())
    ruta_pdf_abs = str(Path(ruta_pdf).resolve())

    # 17 = wdFormatPDF
    WD_FORMAT_PDF = 17

    pythoncom = __import__("pythoncom")
    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(ruta_docx_abs, ReadOnly=True)
        doc.SaveAs(ruta_pdf_abs, FileFormat=WD_FORMAT_PDF)
    except pywintypes.com_error as e:
        raise RuntimeError(f"Error convirtiendo a PDF con Word: {e}")
    finally:
        if doc is not None:
            doc.Close(False)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()


def _crear_pdf_word_linux(ruta_docx: str, ruta_pdf: str) -> None:
    """Convierte un .docx a .pdf usando LibreOffice en modo headless (Linux/VPS)."""
    carpeta_salida = str(Path(ruta_pdf).parent)
    resultado = subprocess.run(
        [
            "libreoffice", "--headless", "--norestore",
            "--convert-to", "pdf",
            "--outdir", carpeta_salida,
            ruta_docx,
        ],
        capture_output=True,
        text=True,
        timeout=60,
    )
    if resultado.returncode != 0:
        raise RuntimeError(f"Error convirtiendo a PDF: {resultado.stderr}")

    # LibreOffice nombra el PDF igual que el .docx (mismo nombre base).
    # Si ruta_pdf pide un nombre distinto, lo renombramos.
    nombre_generado = Path(carpeta_salida) / (Path(ruta_docx).stem + ".pdf")
    if str(nombre_generado) != ruta_pdf and nombre_generado.exists():
        nombre_generado.rename(ruta_pdf)


def crear_pdf_word(ruta_docx: str, ruta_pdf: str) -> None:
    """
    Convierte un .docx a .pdf.
    - En Windows: usa Microsoft Word vía win32com (requiere Word instalado).
    - En Linux (VPS): usa LibreOffice headless.
    Detecta el SO automáticamente para que el mismo código sirva tanto
    en la PC local (Windows) como en el VPS (Linux).
    """
    if platform.system() == "Windows":
        _crear_pdf_word_windows(ruta_docx, ruta_pdf)
    else:
        _crear_pdf_word_linux(ruta_docx, ruta_pdf)

def unir_pdfs(lista_pdfs: Iterable[str], salida: str) -> None:
    writer = PdfWriter()

    for ruta in lista_pdfs:
        print("UNIENDO PDF:", ruta)

        lector = PdfReader(ruta)

        for pagina in lector.pages:
            writer.add_page(pagina)

    with open(salida, "wb") as f:
        writer.write(f)


def guardar_pdfs_extra(archivos: List[UploadFile]) -> List[str]:
    """Guarda temporalmente los PDFs adicionales subidos desde el formulario.
    Se respeta el orden en que llegan (el usuario ya los ordenó en el TSX
    con las flechas ↑ ↓ antes de darle a Generar)."""
    rutas: List[str] = []
    tmp_dir = CARPETA_SALIDA / "_tmp_extra"
    tmp_dir.mkdir(exist_ok=True)
    for i, archivo in enumerate(archivos):
        if not archivo.filename or not archivo.filename.lower().endswith(".pdf"):
            continue
        destino = tmp_dir / f"{i}_{archivo.filename}"
        with open(destino, "wb") as f:
            shutil.copyfileobj(archivo.file, f)
        rutas.append(str(destino))
    return rutas


def limpiar_temporales(rutas: Iterable[str]) -> None:
    for r in rutas:
        try:
            os.remove(r)
        except OSError:
            pass